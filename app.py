from __future__ import annotations

import base64
import io
import json
import os
import re
from pathlib import Path
from typing import Any

from flask import Flask, jsonify, render_template, request, send_file
from docx import Document
from docx.shared import Inches, Pt
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas as pdf_canvas

BASE_DIR = Path(__file__).resolve().parent
DATA_FILE = BASE_DIR / "data" / "invitation_templates.json"
STATIC_CSS_FILE = BASE_DIR / "static" / "styles.css"
STATIC_JS_FILE = BASE_DIR / "static" / "app.js"

app = Flask(__name__)


def load_template_payload() -> dict[str, Any]:
    with DATA_FILE.open("r", encoding="utf-8") as template_file:
        return json.load(template_file)


def build_assets_version() -> str:
    css_version = int(STATIC_CSS_FILE.stat().st_mtime) if STATIC_CSS_FILE.exists() else 0
    js_version = int(STATIC_JS_FILE.stat().st_mtime) if STATIC_JS_FILE.exists() else 0
    return str(max(css_version, js_version))


def sanitize_filename(raw_name: str) -> str:
    safe = re.sub(r"[^a-zA-Z0-9._-]+", "-", raw_name).strip("-._")
    return safe or "undangan"


def decode_data_url_image(data_url: str) -> bytes:
    if not isinstance(data_url, str) or "," not in data_url:
        raise ValueError("Format image_data tidak valid.")

    header, encoded = data_url.split(",", 1)
    if ";base64" not in header:
        raise ValueError("image_data harus berupa data URL base64.")

    try:
        return base64.b64decode(encoded)
    except Exception as exc:  # noqa: BLE001
        raise ValueError("Gagal decode image_data base64.") from exc


def parse_positive_float(value: Any, fallback: float) -> float:
    try:
        parsed = float(value)
    except (TypeError, ValueError):
        return fallback
    return max(100.0, parsed)


def sanitize_text(value: Any) -> str:
    if value is None:
        return ""
    return str(value)


def iter_sorted_components(draft: dict[str, Any]) -> list[dict[str, Any]]:
    components = draft.get("components", [])
    if not isinstance(components, list):
        return []
    return sorted(
        [component for component in components if isinstance(component, dict)],
        key=lambda item: (int(item.get("y", 0)), int(item.get("x", 0))),
    )


def generate_docx_from_draft(draft: dict[str, Any]) -> io.BytesIO:
    document = Document()
    paper = draft.get("paper", {})
    width_px = int(paper.get("width", 900))
    height_px = int(paper.get("height", 1240))

    section = document.sections[0]
    section.page_width = Inches(max(4.0, width_px / 96))
    section.page_height = Inches(max(4.0, height_px / 96))
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    title = sanitize_text(draft.get("name", "Draft Undangan"))
    document.add_heading(title, level=1)
    intro = document.add_paragraph("Dokumen ini dibuat otomatis dari Studio Undangan Nusantara.")
    intro.runs[0].font.size = Pt(10)

    body_types = {"text", "multiline"}
    image_types = {"image", "signature"}
    content_width_in = max(2.5, (width_px / 96) - 1.0)

    for component in iter_sorted_components(draft):
        component_type = sanitize_text(component.get("type", "text")).lower()
        label = sanitize_text(component.get("label", component.get("id", "Komponen")))
        value = sanitize_text(component.get("value", ""))

        if component_type in body_types:
            paragraph = document.add_paragraph()
            label_run = paragraph.add_run(f"{label}: ")
            label_run.bold = True
            paragraph.add_run(value if value else "(kosong)")
            continue

        if component_type in image_types:
            heading = document.add_paragraph()
            heading_run = heading.add_run(f"{label}:")
            heading_run.bold = True

            if value.startswith("data:image"):
                try:
                    image_bytes = decode_data_url_image(value)
                    img_buffer = io.BytesIO(image_bytes)
                    comp_width_px = int(component.get("width", 220))
                    width_in = max(0.8, min(content_width_in, comp_width_px / 96))
                    document.add_picture(img_buffer, width=Inches(width_in))
                except ValueError:
                    document.add_paragraph("(gambar tidak valid)")
            elif value:
                document.add_paragraph(f"(sumber gambar eksternal) {value}")
            else:
                document.add_paragraph("(gambar belum diisi)")

    output = io.BytesIO()
    document.save(output)
    output.seek(0)
    return output


@app.route("/")
def index() -> str:
    payload = load_template_payload()
    return render_template(
        "index.html",
        app_name=payload.get("app_name", "Studio Undangan"),
        asset_v=build_assets_version(),
    )


@app.route("/api/templates")
def templates() -> Any:
    return jsonify(load_template_payload())


@app.route("/healthz")
def healthz() -> Any:
    return jsonify({"ok": True, "service": "studio-undangan"}), 200


@app.route("/api/export/pdf", methods=["POST"])
def export_pdf() -> Any:
    payload = request.get_json(silent=True) or {}
    image_data = payload.get("image_data")
    filename = sanitize_filename(str(payload.get("filename", "undangan")))

    try:
        image_bytes = decode_data_url_image(image_data)
        image_reader = ImageReader(io.BytesIO(image_bytes))
        image_width, image_height = image_reader.getSize()
    except ValueError as err:
        return jsonify({"ok": False, "error": str(err)}), 400
    except Exception:  # noqa: BLE001
        return jsonify({"ok": False, "error": "Gagal membaca gambar untuk PDF."}), 400

    width_pt = parse_positive_float(payload.get("page_width_pt"), image_width * 0.75)
    height_pt = parse_positive_float(payload.get("page_height_pt"), image_height * 0.75)

    output = io.BytesIO()
    pdf = pdf_canvas.Canvas(output, pagesize=(width_pt, height_pt))
    scale = min(width_pt / image_width, height_pt / image_height)
    draw_w = image_width * scale
    draw_h = image_height * scale
    draw_x = (width_pt - draw_w) / 2
    draw_y = (height_pt - draw_h) / 2
    pdf.drawImage(image_reader, draw_x, draw_y, width=draw_w, height=draw_h, preserveAspectRatio=True, mask="auto")
    pdf.showPage()
    pdf.save()
    output.seek(0)

    return send_file(
        output,
        mimetype="application/pdf",
        as_attachment=True,
        download_name=f"{filename}.pdf",
        max_age=0,
    )


@app.route("/api/export/docx", methods=["POST"])
def export_docx() -> Any:
    payload = request.get_json(silent=True) or {}
    filename = sanitize_filename(str(payload.get("filename", "undangan")))
    draft = payload.get("draft")

    if not isinstance(draft, dict):
        return jsonify({"ok": False, "error": "Payload draft tidak valid."}), 400

    if not isinstance(draft.get("components"), list):
        return jsonify({"ok": False, "error": "Komponen draft tidak ditemukan."}), 400

    try:
        output = generate_docx_from_draft(draft)
    except Exception:  # noqa: BLE001
        return jsonify({"ok": False, "error": "Gagal membuat file DOCX."}), 400

    return send_file(
        output,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=f"{filename}.docx",
        max_age=0,
    )


if __name__ == "__main__":
    app.run(
        host="0.0.0.0",
        port=int(os.getenv("PORT", "5000")),
        debug=os.getenv("FLASK_DEBUG", "1") == "1",
    )
